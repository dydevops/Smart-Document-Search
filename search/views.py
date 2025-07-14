import os
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
from django.shortcuts import render, redirect
from .models import Document
from .forms import DocumentUploadForm
from .tasks import index_document_task
# from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangDoc
import re


# def highlight_terms(text, query):
#     for word in query.split():
#         text = re.sub(f"(?i)({re.escape(word)})", r"<strong>\1</strong>", text)
#     return text
# Load once globally
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
try:
    vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
except:
    vectorstore = None



# def extract_text(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext in ['.jpg', '.png']:
#         return pytesseract.image_to_string(Image.open(file_path))
#     elif ext == '.pdf':
#         doc = fitz.open(file_path)
#         return "\n".join(page.get_text() for page in doc)
#     elif ext == '.docx':
#         from docx import Document as DocxDocument
#         doc = DocxDocument(file_path)
#         return "\n".join(p.text for p in doc.paragraphs)
#     return ""



# def upload_document(request):
#     if request.method == 'POST':
#         form = DocumentUploadForm(request.POST, request.FILES)
#         if form.is_valid():
#             doc = form.save(commit=False)
#             doc.user = request.user
#             doc.save()
#             text = extract_text(doc.file.path)
#             save_to_faiss(text, doc.title)
#             return redirect('search')
#     else:
#         form = DocumentUploadForm()
#     return render(request, 'search/upload.html', {'form': form})


# def save_to_faiss(text, title):
#     splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#     chunks = splitter.split_text(text)
#     docs = [LangDoc(page_content=chunk, metadata={"source": title}) for chunk in chunks]
    
#     embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

#     if not os.path.exists("faiss_index"):
#         vectorstore = FAISS.from_documents(docs, embeddings)
#     else:
#         vectorstore = FAISS.load_local("faiss_index", embeddings)
#         vectorstore.add_documents(docs)

#     vectorstore.save_local("faiss_index")


# # Load the model once
# embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# # Try to load FAISS index once
# try:
#     vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
# except:
#     vectorstore = None


# def search(request):
#     results = []
#     query = ""
#     if request.method == 'POST':
#         query = request.POST.get('query')
#         embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
#         vectorstore = FAISS.load_local(
#             "faiss_index", 
#             embeddings, 
#             allow_dangerous_deserialization=True  # ✅ Add this
#         )
#         docs = vectorstore.similarity_search(query, k=5)
#         results = [doc.page_content for doc in docs]

#     return render(request, 'search/search.html', {'results': results, 'query': query})

# def search(request):
#     results = []
#     query = ""
#     if request.method == 'POST':
#         query = request.POST.get('query')
#         if vectorstore:
#             docs = vectorstore.similarity_search(query, k=3)  # use k=3 for faster response
#             results = [doc.page_content for doc in docs]
#     return render(request, 'search/search.html', {'results': results, 'query': query})
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.jpg', '.png']:
        return pytesseract.image_to_string(Image.open(file_path))
    elif ext == '.pdf':
        doc = fitz.open(file_path)
        return "\n".join(page.get_text() for page in doc)
    elif ext == '.docx':
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    return ""

def upload_document(request):
    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            doc = form.save(commit=False)
            doc.user = request.user
            doc.save()
            text = extract_text(doc.file.path)
            index_document_task.delay(text, doc.title)  # Celery async
            return redirect('search')
    else:
        form = DocumentUploadForm()
    return render(request, 'search/upload.html', {'form': form})

def search(request):
    results = []
    query = ""
    if request.method == 'POST':
        query = request.POST.get('query')
        if vectorstore:
            docs = vectorstore.similarity_search(query, k=5)
            results = [doc.page_content for doc in docs]
    return render(request, 'search/search.html', {'results': results, 'query': query})



# def semantic_search(request):
#     query = None
#     results = []

#     if request.method == 'POST':
#         query = request.POST.get('query', '')
#         if query:
#             embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
#             vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
#             docs = vectorstore.similarity_search(query, k=4)
#             results = [highlight_terms(doc.page_content, query) for doc in docs]

#     return render(request, "search/semantic_search.html", {
#         "query": query,
#         "results": results
#     })

def highlight_terms(text, query):
    """
    Highlights all instances of the query in the text using <mark> tag.
    """
    pattern = re.compile(re.escape(query), re.IGNORECASE)
    return pattern.sub(lambda m: f'<mark>{m.group(0)}</mark>', text)

def semantic_search(request):
    query = ""
    results = []

    if request.method == 'POST':
        query = request.POST.get('query', '').strip()
        if query:
            embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
            vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
            docs = vectorstore.similarity_search(query, k=4)

            # Only include docs with similarity hits (basic filter)
            results = [highlight_terms(doc.page_content, query) for doc in docs]

    return render(request, "search/semantic_search.html", {
        "query": query,
        "results": results
    })